#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""上市公司公告案例分析报告 -> Word (.docx) 生成器（美化版 v2）

用法:
    build_docx.py <input.json> <output.docx>

特性:
- 封面页（主题色块 + 元信息）
- 页眉（标题+主题）/ 页脚（页码）
- 结论速览 / 法律分析 / 局限性 三色块卡片
- 案例卡片化（深色标题栏 + 浅底内容区）
- 案例索引表（深色表头 + 斑马纹）
- 字体跨平台回退（Win YaHei / macOS PingFang / Linux Noto）
- 正文首行缩进 + 1.5 倍行距

输入 JSON 结构见 references/report_template.md。
依赖: python-docx
"""
import sys
import json
import platform
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- 配色 ----
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0xDD, 0xE6, 0xF0)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)   # 结论强调色（红）
GREEN = RGBColor(0x1E, 0x7E, 0x34)    # 达标绿
WARN = RGBColor(0xB9, 0x6A, 0x00)      # 局限性橙
BOX_BLUE = "EAF1FB"      # 速览卡片浅蓝底
BOX_GREEN = "EDF5ED"     # 法律分析浅绿底
BOX_ORANGE = "FDF3E7"    # 局限性浅橙底
CARD_BODY = "F7F9FC"     # 案例内容区浅底
ZEBRA = "F2F6FB"         # 索引表斑马纹
BORDER = "C9D6E5"        # 浅边框色


def get_cjk_font():
    s = platform.system()
    if s == "Windows":
        return "Microsoft YaHei"
    if s == "Darwin":
        return "PingFang SC"
    return "Noto Sans CJK SC"
CJK = get_cjk_font()


def set_run_font(run, size=None, bold=None, color=None, font=CJK):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.append(rf)
    rf.set(qn('w:eastAsia'), font)
    rf.set(qn('w:ascii'), font)
    rf.set(qn('w:hAnsi'), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    rid = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyper = OxmlElement('w:hyperlink')
    hyper.set(qn('r:id'), rid)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    hyper.append(r)
    paragraph._p.append(hyper)
    for rr in hyper.findall(qn('w:r')):
        rpr = rr.get_or_add_rPr()
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rpr.append(color)
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rpr.append(u)


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def set_cell_border(cell, color='1F3A5F', sz='10'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tcPr.append(borders)


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right)):
        e = OxmlElement('w:' + tag)
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')
        m.append(e)
    tcPr.append(m)


def cell_para(cell, text="", size=10.5, bold=False, color=None, space_after=3,
              first=False, indent=0, line=1.4, icon=None, icon_color=None):
    """在单元格内加一段文字（first=True 用第 0 段，否则新增段）。"""
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    p.paragraph_format.line_spacing = line
    if icon:
        set_run_font(p.add_run(icon + " "), size=size, bold=True,
                     color=icon_color or color or NAVY)
    if text:
        set_run_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def kv_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(label + "："), size=10.5, bold=True, color=GREY)
    set_run_font(p.add_run(value or "—"), size=10.5)
    return p


def section(doc, title, space_before=10):
    """标题：左侧竖向色条 + 加粗 navy 标题。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '5')
    left.set(qn('w:color'), '1F3A5F')
    pbdr.append(left)
    pPr.append(pbdr)
    set_run_font(p.add_run("  " + title), size=12.5, bold=True, color=NAVY)
    return p


def para(doc, text, size=10.5, bold=False, color=None, space=4, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)  # 首行缩进 2 字符
    set_run_font(p.add_run(text or "—"), size=size, bold=bold, color=color)
    return p


def para_raw(doc, text, size=10.5, bold=False, color=None, space=4, indent=18):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.left_indent = Pt(indent)
    set_run_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def _box_line(cell, icon, label, value, first=False,
              icon_color=NAVY, label_color=NAVY, value_color=None,
              value_bold=False, value_size=10.5):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    if icon:
        set_run_font(p.add_run(icon + " "), size=value_size, bold=True, color=icon_color)
    if label:
        set_run_font(p.add_run(label + "："), size=value_size, bold=True, color=label_color)
    if value:
        set_run_font(p.add_run(value), size=value_size, bold=value_bold, color=value_color)
    return p


def summary_box(doc, data, cases):
    """文档最开头的『结论速览』卡片（第一色块）。"""
    meta = data.get('meta', {})
    n = len(cases)
    verdict = data.get('verdict') or ''
    if not verdict:
        concl = data.get('conclusion', '') or data.get('overview', '')
        verdict = (concl.split('。')[0] + '。') if '。' in concl else concl[:140]
    findings = (data.get('core_findings') or [])[:3]

    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(4)
    tp.paragraph_format.space_after = Pt(2)
    set_run_font(tp.add_run("■ 结论速览  EXECUTIVE SUMMARY"), size=12.5, bold=True, color=NAVY)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, BOX_BLUE)
    set_cell_border(cell, color='1F3A5F', sz='8')
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)

    first = True
    if meta.get('topic'):
        _box_line(cell, "◆", "主题", meta.get('topic'), first=True)
        first = False

    if n >= 20:
        cnt_txt = "%d 个（满足 ≥20 案例下限）" % n
        cnt_color = GREEN
    else:
        cnt_txt = "%d 个（低于建议下限 20，详见局限性）" % n
        cnt_color = WARN
    _box_line(cell, "●", "案例数", cnt_txt, first=first, value_color=cnt_color, value_bold=True)

    _box_line(cell, "★", "核心结论", verdict, value_color=ACCENT, value_bold=True, value_size=11)

    for f in findings:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.line_spacing = 1.2
        set_run_font(p.add_run("→ "), size=10, bold=True, color=NAVY)
        set_run_font(p.add_run(f), size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def color_block(doc, title, title_en, body_lines, bg_hex,
                title_color=NAVY, body_color=GREY, icon=None):
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(4)
    tp.paragraph_format.space_after = Pt(2)
    set_run_font(tp.add_run("■ %s  %s" % (title, title_en)), size=12.5, bold=True, color=title_color)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg_hex)
    set_cell_border(cell, color='1F3A5F', sz='8')
    set_cell_margins(cell, top=90, bottom=90, left=140, right=140)
    first = True
    for ln in body_lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.4
        if icon:
            set_run_font(p.add_run(icon + " "), size=10, bold=True, color=title_color)
        set_run_font(p.add_run(ln), size=10, color=body_color)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def legal_block(doc, data):
    legal = data.get('legal_analysis') or ''
    if legal:
        color_block(doc, "法律分析", "LEGAL ANALYSIS", [legal],
                    bg_hex=BOX_GREEN, title_color=GREEN, body_color=GREY)


def limitations_block(doc, data):
    lms = data.get('limitations') or []
    if not lms and data.get('note'):
        lms = [data.get('note')]
    if lms:
        color_block(doc, "局限性", "LIMITATIONS", lms,
                    bg_hex=BOX_ORANGE, title_color=WARN, body_color=GREY, icon="※")


def cover_page(doc, meta):
    """封面：顶部色块大标题 + 副标题 + 元信息块。"""
    doc.add_paragraph().paragraph_format.space_after = Pt(72)
    # 主题色块
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "1F3A5F")
    set_cell_border(cell, color='1F3A5F', sz='0')
    set_cell_margins(cell, top=260, bottom=260, left=200, right=200)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("上市公司公告案例分析报告"), size=22, bold=True, color=WHITE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(6)
    set_run_font(p2.add_run(meta.get('topic', '') or '—'), size=13, bold=True, color=LIGHT)

    doc.add_paragraph().paragraph_format.space_after = Pt(48)
    # 元信息块（左标签 + 右值，居中）
    info = [
        ("检索范围", meta.get('scope', '')),
        ("案例类型", meta.get('case_type', '')),
        ("数据来源", meta.get('mode', '')),
        ("生成日期", meta.get('generated_at', '')),
    ]
    itbl = doc.add_table(rows=len(info), cols=2)
    itbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    itbl.columns[0].width = Inches(1.4)
    itbl.columns[1].width = Inches(4.6)
    for i, (k, v) in enumerate(info):
        c0 = itbl.rows[i].cells[0]
        c1 = itbl.rows[i].cells[1]
        set_cell_border(c0, sz='0')
        set_cell_border(c1, sz='0')
        c0.width = Inches(1.4)
        c1.width = Inches(4.6)
        set_run_font(c0.paragraphs[0].add_run(k), size=10.5, bold=True, color=NAVY)
        set_run_font(c1.paragraphs[0].add_run(v or '—'), size=10.5, color=GREY)
    doc.add_page_break()


def set_header_footer(section_obj, title, subtitle):
    # 页眉：左标题 + 右主题
    header = section_obj.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(hp.add_run(title), size=9, color=NAVY, bold=True)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)
    run = hp.add_run("\t" + (subtitle or ""))
    set_run_font(run, size=9, color=GREY)
    # 页脚：页码
    footer = section_obj.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(fp.add_run("第 "), size=9, color=GREY)
    fld1 = OxmlElement('w:fldSimple')
    fld1.set(qn('w:instr'), 'PAGE')
    r1 = OxmlElement('w:r')
    t1 = OxmlElement('w:t')
    t1.text = '1'
    r1.append(t1)
    fld1.append(r1)
    fp._p.append(fld1)
    set_run_font(fp.add_run(" 页 / 共 "), size=9, color=GREY)
    fld2 = OxmlElement('w:fldSimple')
    fld2.set(qn('w:instr'), 'NUMPAGES')
    r2 = OxmlElement('w:r')
    t2 = OxmlElement('w:t')
    t2.text = '1'
    r2.append(t2)
    fld2.append(r2)
    fp._p.append(fld2)
    set_run_font(fp.add_run(" 页"), size=9, color=GREY)


def index_table(doc, cases):
    tbl = doc.add_table(rows=1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    heads = ["#", "公司", "代码", "公告标题", "日期"]
    widths = [Inches(0.4), Inches(1.5), Inches(0.8), Inches(3.4), Inches(1.0)]
    for i, h in enumerate(heads):
        set_cell_bg(hdr[i], "1F3A5F")
        set_cell_border(hdr[i], color='1F3A5F', sz='6')
        hdr[i].width = widths[i]
        set_run_font(hdr[i].paragraphs[0].add_run(h), size=9.5, bold=True, color=WHITE)
    for idx, c in enumerate(cases, 1):
        row = tbl.add_row().cells
        vals = [str(idx), c.get('company', ''), c.get('code', ''),
                c.get('title', ''), c.get('date', '')]
        for i, v in enumerate(vals):
            row[i].width = widths[i]
            set_run_font(row[i].paragraphs[0].add_run(v), size=9, color=GREY)
            if idx % 2 == 0:
                set_cell_bg(row[i], ZEBRA)
    return tbl


def case_card(doc, idx, c):
    """单个案例：深蓝标题栏 + 浅底内容区卡片。"""
    # 标题栏
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "1F3A5F")
    set_cell_border(cell, color='1F3A5F', sz='0')
    set_cell_margins(cell, top=70, bottom=70, left=140, right=140)
    p = cell.paragraphs[0]
    set_run_font(p.add_run("案例 %d" % idx), size=11, bold=True, color=WHITE)
    p.add_run("    ").font.size = Pt(11)
    set_run_font(p.add_run("%s（%s）" % (c.get('company', ''), c.get('code', ''))),
                 size=11, bold=True, color=WHITE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    set_run_font(p2.add_run(c.get('title', '')), size=9.5, color=LIGHT)

    # 内容区
    body = doc.add_table(rows=1, cols=1)
    body.alignment = WD_TABLE_ALIGNMENT.CENTER
    bcell = body.rows[0].cells[0]
    set_cell_bg(bcell, CARD_BODY)
    set_cell_border(bcell, color=BORDER, sz='6')
    set_cell_margins(bcell, top=110, bottom=110, left=160, right=160)

    cell_para(bcell, "交易所：%s　|　日期：%s　|　类型：%s" % (
        c.get('exchange', ''), c.get('date', ''), c.get('type', '')),
        size=9.5, bold=True, color=NAVY, space_after=6, first=True)
    cell_para(bcell, "【案例概要】", size=10.5, bold=True, color=NAVY, space_after=2)
    cell_para(bcell, c.get('summary', ''), size=10.5, space_after=6, indent=18, line=1.5)
    cell_para(bcell, "【案例分析】", size=10.5, bold=True, color=NAVY, space_after=2)
    cell_para(bcell, c.get('analysis', ''), size=10.5, space_after=6, indent=18, line=1.5)
    if c.get('regulatory'):
        cell_para(bcell, "【监管视角 / 可借鉴点】", size=10.5, bold=True, color=NAVY, space_after=2)
        cell_para(bcell, c.get('regulatory', ''), size=10.5, space_after=6, indent=18, line=1.5)
    quotes = c.get('quotes') or []
    if quotes:
        cell_para(bcell, "【引用片段】", size=10.5, bold=True, color=NAVY, space_after=2)
        for q in quotes:
            cell_para(bcell, "“%s”" % q, size=9.5, color=GREY, space_after=3, indent=18, line=1.4)
    if c.get('link'):
        p = bcell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run("原文链接："), size=10.5, bold=True, color=GREY)
        add_hyperlink(p, c['link'], c['link'])
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def main():
    if len(sys.argv) < 3:
        print("usage: build_docx.py <input.json> <output.docx>")
        sys.exit(1)
    inp, out = sys.argv[1], sys.argv[2]
    with open(inp, encoding='utf-8') as f:
        data = json.load(f)

    meta = data.get('meta', {})
    cases = data.get('cases', [])
    conclusion = data.get('conclusion', '')
    notes = data.get('notes', '')

    doc = Document()
    # 页边距
    for s in doc.sections:
        s.top_margin = Inches(0.9)
        s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(0.95)
        s.right_margin = Inches(0.95)
        set_header_footer(s, "上市公司公告案例分析报告", meta.get('topic', ''))

    nstyle = doc.styles['Normal']
    rpr = nstyle.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.append(rf)
    rf.set(qn('w:eastAsia'), CJK)
    rf.set(qn('w:ascii'), CJK)
    rf.set(qn('w:hAnsi'), CJK)

    # 封面
    cover_page(doc, meta)

    # 速览 / 法律 / 局限性 三色块
    summary_box(doc, data, cases)
    legal_block(doc, data)
    limitations_block(doc, data)

    # 总览
    if data.get('overview'):
        section(doc, "总览")
        para(doc, data['overview'], space=6)

    # 核心发现
    cf = data.get('core_findings') or []
    if cf:
        section(doc, "核心发现")
        for item in cf:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.line_spacing = 1.4
            set_run_font(p.add_run("➤ "), size=10.5, bold=True, color=NAVY)
            set_run_font(p.add_run(item), size=10.5)

    # 案例索引
    section(doc, "一、案例索引")
    if cases:
        index_table(doc, cases)

    # 逐案例卡片
    for idx, c in enumerate(cases, 1):
        case_card(doc, idx, c)

    if conclusion:
        section(doc, "二、总体结论")
        para(doc, conclusion, space=6)
    if notes:
        section(doc, "附注")
        para_raw(doc, notes, size=9.5, color=GREY, space=4, indent=0)

    doc.save(out)
    print("OK", out, "cases=", len(cases))


if __name__ == '__main__':
    main()
